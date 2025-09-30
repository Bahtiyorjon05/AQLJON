import asyncio
import tempfile
import logging
from telegram import Update
from telegram.ext import ContextTypes
from telegram.constants import ParseMode
from modules.doc_generation.base_generator import BaseDocumentGenerator
from modules.config import Config
from modules.utils import safe_reply, send_typing
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_TABLE_ALIGNMENT
import random

logger = logging.getLogger(__name__)

class WordGenerator(BaseDocumentGenerator):
    """Handles Word document generation with enhanced styling and professional themes.
    
    This class generates professional Word documents with advanced styling,
    Unicode support, and theme-based color schemes. It supports various document
    types and automatically selects appropriate fonts for better emoji and 
    special character display.
    """
    
    def _get_current_date(self):
        """Get current date in a formatted string"""
        from datetime import datetime
        return datetime.now().strftime("%d-%m-%Y")
    
    def _get_theme_colors(self, theme_name='modern'):
        """Get a comprehensive color palette for a specific theme with enhanced color options"""
        themes = {
            'modern': {
                'primary': '#2C3E50',      # Deep navy
                'secondary': '#3498DB',    # Bright blue
                'accent': '#E74C3C',       # Vibrant red
                'background': '#ECF0F1',   # Light gray
                'text': '#2C3E50',         # Dark navy
                'light_text': '#7F8C8D',   # Gray
                'highlight': '#F39C12',    # Amber
                'success': '#27AE60',      # Green
                'warning': '#F39C12',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#BDC3C7',      # Light gray-blue
                'education': '#9B59B6',    # Purple for educational content
                'insight': '#1ABC9C',      # Turquoise for insights
                'quote': '#8E44AD',        # Deep purple for quotes
                'code': '#F39C12',         # Amber for code blocks
                'link': '#3498DB',         # Blue for links
            },
            'elegant': {
                'primary': '#8E44AD',      # Purple
                'secondary': '#9B59B6',    # Light purple
                'accent': '#E67E22',       # Orange
                'background': '#F8F9F9',   # Very light gray
                'text': '#2C3E50',         # Dark navy
                'light_text': '#95A5A6',   # Gray
                'highlight': '#1ABC9C',    # Turquoise
                'success': '#27AE60',      # Green
                'warning': '#F39C12',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#D5DBDB',      # Light gray
                'education': '#8E44AD',    # Purple for educational content
                'insight': '#1ABC9C',      # Turquoise for insights
                'quote': '#9B59B6',        # Light purple for quotes
                'code': '#E67E22',         # Orange for code blocks
                'link': '#9B59B6',         # Light purple for links
            },
            'vibrant': {
                'primary': '#E74C3C',      # Red
                'secondary': '#E67E22',    # Orange
                'accent': '#F1C40F',       # Yellow
                'background': '#FDEDEC',   # Light red background
                'text': '#2C3E50',         # Dark navy
                'light_text': '#95A5A6',   # Gray
                'highlight': '#1ABC9C',    # Turquoise
                'success': '#27AE60',      # Green
                'warning': '#F39C12',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#EBEDEF',      # Light gray
                'education': '#E67E22',    # Orange for educational content
                'insight': '#1ABC9C',      # Turquoise for insights
                'quote': '#E74C3C',        # Red for quotes
                'code': '#F1C40F',         # Yellow for code blocks
                'link': '#E67E22',         # Orange for links
            },
            'professional': {
                'primary': '#1F3A93',      # Deep blue
                'secondary': '#3498DB',    # Bright blue
                'accent': '#26A65B',       # Green
                'background': '#F8F9F9',   # Very light gray
                'text': '#2C3E50',         # Dark navy
                'light_text': '#7F8C8D',   # Gray
                'highlight': '#F39C12',    # Amber
                'success': '#27AE60',      # Green
                'warning': '#F39C12',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#D5DBDB',      # Light gray
                'education': '#1F3A93',    # Deep blue for educational content
                'insight': '#26A65B',      # Green for insights
                'quote': '#1F3A93',        # Deep blue for quotes
                'code': '#F39C12',         # Amber for code blocks
                'link': '#3498DB',         # Blue for links
            },
            'corporate': {
                'primary': '#2C3E50',      # Charcoal
                'secondary': '#34495E',    # Dark blue-gray
                'accent': '#16A085',       # Teal
                'background': '#FFFFFF',   # White
                'text': '#2C3E50',         # Charcoal
                'light_text': '#95A5A6',   # Gray
                'highlight': '#F39C12',    # Amber
                'success': '#27AE60',      # Green
                'warning': '#F39C12',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#ECF0F1',      # Light gray
                'education': '#2C3E50',    # Charcoal for educational content
                'insight': '#16A085',      # Teal for insights
                'quote': '#34495E',        # Dark blue-gray for quotes
                'code': '#16A085',         # Teal for code blocks
                'link': '#3498DB',         # Blue for links
            },
            'creative': {
                'primary': '#9B59B6',      # Purple
                'secondary': '#3498DB',    # Blue
                'accent': '#E74C3C',       # Red
                'background': '#F9F5FF',   # Light purple background
                'text': '#2C3E50',         # Dark navy
                'light_text': '#7F8C8D',   # Gray
                'highlight': '#F1C40F',    # Yellow
                'success': '#27AE60',      # Green
                'warning': '#E67E22',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#D5DBDB',      # Light gray
                'education': '#9B59B6',    # Purple for educational content
                'insight': '#1ABC9C',      # Turquoise for insights
                'quote': '#9B59B6',        # Purple for quotes
                'code': '#F1C40F',         # Yellow for code blocks
                'link': '#9B59B6',         # Purple for links
            },
            'minimalist': {
                'primary': '#2C3E50',      # Dark navy
                'secondary': '#7F8C8D',    # Gray
                'accent': '#3498DB',       # Blue
                'background': '#FFFFFF',   # White
                'text': '#2C3E50',         # Dark navy
                'light_text': '#BDC3C7',   # Light gray
                'highlight': '#3498DB',    # Blue
                'success': '#27AE60',      # Green
                'warning': '#F39C12',      # Orange
                'info': '#3498DB',         # Blue
                'divider': '#ECF0F1',      # Light gray
                'education': '#2C3E50',    # Dark navy for educational content
                'insight': '#7F8C8D',      # Gray for insights
                'quote': '#7F8C8D',        # Gray for quotes
                'code': '#3498DB',         # Blue for code blocks
                'link': '#3498DB',         # Blue for links
            }
        }
        return themes.get(theme_name, themes['modern'])
    
    async def generate(self, update: Update, context: ContextTypes.DEFAULT_TYPE, topic: str, content_context: str = ""):
        """Generate a professional Word document based on user request"""
        # Validate topic using centralized validation
        is_valid, cleaned_topic = self._validate_topic(topic)
        if not is_valid:
            await self._track_document_generation(update, "word")
            if len(cleaned_topic) < 2:
                processing_msg = await self._send_processing_message(update, "<b>❌ Word hujjat yaratishda xatolik.</b>\n\nMavzu juda qisqa.")
                if processing_msg:
                    await processing_msg.edit_text("❌ Iltimos, Word hujjati uchun to'liq mavzu kiriting (kamida 2 belgi).", parse_mode=ParseMode.HTML)
            else:
                processing_msg = await self._send_processing_message(update, "<b>❌ Word hujjat yaratishda xatolik.</b>\n\nIltimos, mavzu kiriting.")
                if processing_msg:
                    await processing_msg.edit_text("❌ Iltimos, Word hujjati uchun mavzu kiriting.", parse_mode=ParseMode.HTML)
            return
        
        # Track document generation
        await self._track_document_generation(update, "word")
        
        # Initialize processing_msg to avoid unbound variable error
        processing_msg = None
        
        # Send immediate processing message for better user experience
        try:
            from modules.utils import send_fast_reply
            if update.message:
                send_fast_reply(update.message, "<b>📝 Word hujjatni tuzyapman. Biroz kutib turing... ⏳</b>")
                # Send typing indicator
                await send_typing(update)
        except:
            # Fallback if fast reply fails
            processing_msg = await self._send_processing_message(update, f"<b>📝 Word hujjatni tuzyapman. Biroz kutib turing... ⏳</b>")
            # Send typing indicator
            await send_typing(update)
        
        try:
            # Generate filename using centralized method
            filename = await self._generate_filename(cleaned_topic, "Word")
            
            # Generate content with Gemini using flexible prompt structure - all content in user's language
            # Gemini will automatically detect the language from the user's input
            prompt = f"""
            You are AQLJON, a Muslim intelligent assistant who creates exceptional, professionally formatted Word documents.
            Create a professional Word document about '{cleaned_topic}' in the SAME LANGUAGE as the user's input.
            
            {"Use the following context from previous documents the user shared to inform your content:" + content_context if content_context else ""}
            
            Analyze the topic and create a well-structured, engaging document that:
            1. Provides medium-detailed coverage based on the topic type (not too detailed nor too brief)
            2. Uses an appropriate structure for a Word document with proper sections
            3. Incorporates bullet points, numbered lists, examples, and visual elements where appropriate
            4. Provides valuable insights and actionable information
            5. Uses professional formatting with proper hierarchy
            
            IMPORTANT GUIDELINES:
            1. If the topic relates to academic problems (math, physics, chemistry, coding, biology):
               - NEVER provide direct solutions or answers
               - Guide with concepts, understanding, and hints
               - Focus on educational value and understanding
               - Force them to think and try again while giving understanding clearly
               - Include learning takeaways and key concepts
            
            2. For general topics:
               - Provide informative, well-researched content
               - Keep explanations clear and engaging
               - Use appropriate examples and practical insights
               - Include statistics, facts, or data where relevant
               - Add educational lessons and takeaways
            
            3. For all topics:
               - Create medium-detailed content (not too detailed nor small)
               - Include educational lessons, learning takeaways, and key insights
               - Add practical examples and real-world applications
               - Incorporate quotes, statistics, and data points to enhance credibility
               - Use varied sentence structures and vocabulary to maintain engagement
               - Create content that flows naturally and is engaging to read
               - Do not be overly concise - create medium-detailed content with educational value
               - Include actionable insights and practical advice
               - Be respectful of Islamic values and practices
            
            4. Never provide illicit, harmful, or inappropriate content
            5. Format the response with proper headings using # for main title, ## for sections, ### for subsections
            6. Use bullet points (-) and numbered lists where appropriate for better readability
            7. Make the content informative, valuable, and professionally structured
            8. Include practical examples and actionable insights where relevant
            9. Adapt the structure and tone based on the topic - be flexible rather than following a rigid template
            10. Create all content in the SAME LANGUAGE as the user's input: '{cleaned_topic}'
            11. Make the content conversational and human-like, avoiding robotic or technical language
            12. Create an engaging main title that captures the essence of the topic (do not use the exact user input as the title)
            13. Do not include any "Generated by" text or similar robotic phrases like "Document on [topic]" or "Educational Document with Learning Takeaways"
            14. Add emojis and formatting to make the content visually appealing and engaging
            15. Include quotes, statistics, and data points where relevant to enhance credibility
            16. Use varied sentence structures and vocabulary to maintain engagement
            17. Create content that flows naturally and is engaging to read
            18. For the main title (#), create something compelling and professional that represents the topic well
            19. For section titles (##), create clear, descriptive headings that organize the content logically
            20. For subsection titles (###), create specific headings that dive deeper into each section
            21. Use bullet points (-) for lists of related items or key points
            22. Use numbered lists for step-by-step processes or ordered information
            23. Include relevant emojis to make the content visually appealing but not excessive
            24. Add practical examples, case studies, or real-world applications where appropriate
            25. Include data, statistics, or quotes to support key points and enhance credibility
            26. Add educational lessons, learning takeaways, and key insights in each section
            27. Create medium-detailed content with educational value (not brief summaries, not overly detailed)
            28. Include actionable insights and practical advice throughout the document
            29. Do not include any robotic or generic phrases like "Document on [topic]" or "Educational Document with Learning Takeaways"
            30. Dont be robotic and provide content in a way that is easy to read and understand and think like a Muslim person not AI
            """
            
            # Generate content with timeout to prevent blocking
            try:
                response = await asyncio.wait_for(
                    asyncio.to_thread(lambda: self.model.generate_content(prompt)),
                    timeout=Config.PROCESSING_TIMEOUT * 2  # Double timeout for document generation
                )
                content = response.text if response and response.text else f"Error generating content for {cleaned_topic}."
            except asyncio.TimeoutError:
                logger.warning("Word content generation timed out, using fallback content")
                content = f"# {cleaned_topic}\n\nProfessional Word document on {cleaned_topic}.\n\nGenerated by AQLJON."
            
            # Create Word document in temporary file - offload to separate thread to avoid blocking asyncio
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                tmp_path = tmp_file.name
            
            try:
                # Offload Word document creation to a separate thread to avoid blocking the asyncio event loop with timeout
                await asyncio.wait_for(
                    asyncio.get_event_loop().run_in_executor(
                        None, 
                        self._create_word_document, 
                        content, tmp_path, cleaned_topic, filename
                    ),
                    timeout=Config.PROCESSING_TIMEOUT  # Configurable timeout for Word creation
                )
                
                # Send Word document to user
                if update.message:
                    with open(tmp_path, 'rb') as word_file:
                        await update.message.reply_document(
                            document=word_file,
                            filename=f"{filename}.docx",
                            caption=f"📝 <b>'{filename}' mavzusida professional Word hujjat</b>\n📁 Fayl nomi: {filename}.docx",
                            parse_mode=ParseMode.HTML
                        )
                
                # Send success message
                try:
                    from modules.utils import send_fast_reply
                    if update.message:
                        send_fast_reply(update.message, f"✅ <b>'{filename}' nomli Word hujjat muvaffaqiyatli tuzildi va yuborildi!</b>\n📥 Chiroyli dizayn va uslubda tuzilgan faylingizdan zavqlaning!", parse_mode=ParseMode.HTML)
                except:
                    # Fallback if fast reply fails
                    if processing_msg is not None:
                        await processing_msg.edit_text(f"✅ <b>'{filename}' nomli Word hujjat muvaffaqiyatli tuzildi va yuborildi!</b>\n📥 Chiroyli dizayn va uslubda tuzilgan faylingizdan zavqlaning!", parse_mode=ParseMode.HTML)
            
            finally:
                # Clean up temporary file
                try:
                    import os
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp file: {e}")
        
        except Exception as e:
            logger.error(f"Word generation error: {e}")
            # Send error message to user
            try:
                from modules.utils import send_fast_reply
                if update.message:
                    send_fast_reply(update.message, "❌ Word hujjat yaratishda xatolik. Iltimos, keyinroq qayta urinib ko'ring.", parse_mode=ParseMode.HTML)
            except:
                # Fallback if fast reply fails
                if processing_msg is not None:
                    await processing_msg.edit_text("❌ Word hujjat yaratishda xatolik. Iltimos, keyinroq qayta urinib ko'ring.", parse_mode=ParseMode.HTML)
    
    def _create_word_document(self, content, tmp_path, topic, filename):
        """Create Word document in a separate thread to avoid blocking asyncio"""
        try:
            # Create professional Word document with enhanced styling
            doc = DocxDocument()
            
            # Set document properties
            doc.core_properties.title = filename
            doc.core_properties.category = "Generated Document"
            doc.core_properties.comments = f"Professional document on {filename}"
            
            # Select a professional color scheme
            color_scheme = self._get_theme_colors('professional')
            
            # Create custom styles with enhanced visual elements
            self._create_custom_styles(doc, color_scheme)
            
            # Parse content and create document elements with enhanced formatting
            lines = content.split('\n')
            
            # Extract main title from content if available
            main_title = filename  # Default to user's input
            content_lines = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    content_lines.append(line)
                # Handle main title (#)
                elif line.startswith('# ') and not line.startswith('##'):
                    main_title = line[2:]  # Extract title without the # 
                    # Don't add to content_lines as we'll use it as the cover page title
                else:
                    content_lines.append(line)
            
            # Add cover page content with enhanced styling
            self._add_cover_page(doc, main_title, topic, color_scheme)
            
            # Add table of contents with enhanced styling
            self._add_table_of_contents(doc, color_scheme)
            
            # Process content lines with enhanced formatting
            in_list = False
            list_type = None
            list_counter = 1
            
            for line in content_lines:
                line = line.strip()
                if not line:
                    # Add paragraph break
                    if in_list:
                        in_list = False
                        list_type = None
                        list_counter = 1
                    doc.add_paragraph()
                    continue
                
                # Handle main headers (##)
                if line.startswith('## '):
                    if in_list:
                        in_list = False
                        list_type = None
                        list_counter = 1
                    self._add_enhanced_heading(doc, line[3:], level=1, colors=color_scheme)  # Remove ## 
                # Handle subheaders (###)
                elif line.startswith('### '):
                    if in_list:
                        in_list = False
                        list_type = None
                        list_counter = 1
                    self._add_enhanced_heading(doc, line[4:], level=2, colors=color_scheme)  # Remove ### 
                # Handle bullet points
                elif line.startswith('- ') or line.startswith('* '):
                    if not in_list or list_type != 'bullet':
                        if in_list:
                            # End previous list
                            in_list = False
                            list_type = None
                            list_counter = 1
                        list_type = 'bullet'
                        in_list = True
                    bullet_text = line[2:]  # Remove - or * 
                    # Add emoji if not present
                    if not any(emoji in bullet_text for emoji in ['✅', '🔹', '🔸', '⭐', '📌', '💡', '🔍', '📈', '📊']):
                        bullet_text = "🔹 " + bullet_text
                    self._add_enhanced_bullet_point(doc, bullet_text, colors=color_scheme)
                # Handle numbered lists
                elif line[0:2].isdigit() and line[1:3] == '. ':
                    if not in_list or list_type != 'numbered':
                        if in_list:
                            # End previous list
                            in_list = False
                            list_type = None
                            list_counter = 1
                        list_type = 'numbered'
                        in_list = True
                        list_counter = 1
                    numbered_text = line[3:]  # Remove number and period
                    self._add_enhanced_numbered_point(doc, numbered_text, colors=color_scheme)
                # Handle blockquotes
                elif line.startswith('>'):
                    quote_text = line[1:].strip()
                    self._add_enhanced_quote(doc, quote_text, colors=color_scheme)
                # Handle regular paragraphs
                else:
                    if in_list:
                        in_list = False
                        list_type = None
                        list_counter = 1
                    self._add_enhanced_paragraph(doc, line, colors=color_scheme)
            
            # Add summary section with enhanced styling
            self._add_summary_section(doc, color_scheme)
            
            # Save the document
            doc.save(tmp_path)
            
        except Exception as e:
            logger.error(f"Error creating Word document: {e}")
            raise
    
    def _create_custom_styles(self, doc, color_scheme):
        """Create custom styles for the document with enhanced visual elements"""
        # Create title style with background and border
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.base_style = doc.styles['Title']
        title_font = title_style.font
        title_font.size = Pt(32)
        title_font.bold = True
        title_font.color.rgb = RGBColor(
            int(color_scheme['primary'][1:3], 16),
            int(color_scheme['primary'][3:5], 16),
            int(color_scheme['primary'][5:7], 16)
        )
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(25)
        title_style.paragraph_format.space_before = Pt(10)
        
        # Create heading 1 style with enhanced styling
        heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.base_style = doc.styles['Heading 1']
        heading1_font = heading1_style.font
        heading1_font.size = Pt(24)
        heading1_font.bold = True
        heading1_font.color.rgb = RGBColor(
            int(color_scheme['secondary'][1:3], 16),
            int(color_scheme['secondary'][3:5], 16),
            int(color_scheme['secondary'][5:7], 16)
        )
        heading1_style.paragraph_format.space_after = Pt(18)
        heading1_style.paragraph_format.space_before = Pt(20)
        heading1_style.paragraph_format.keep_with_next = True
        
        # Create heading 2 style with enhanced styling
        heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.base_style = doc.styles['Heading 2']
        heading2_font = heading2_style.font
        heading2_font.size = Pt(20)
        heading2_font.bold = True
        heading2_font.color.rgb = RGBColor(
            int(color_scheme['text'][1:3], 16),
            int(color_scheme['text'][3:5], 16),
            int(color_scheme['text'][5:7], 16)
        )
        heading2_style.paragraph_format.space_after = Pt(15)
        heading2_style.paragraph_format.space_before = Pt(18)
        heading2_style.paragraph_format.keep_with_next = True
        
        # Create heading 3 style with enhanced styling
        heading3_style = doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        heading3_style.base_style = doc.styles['Heading 3']
        heading3_font = heading3_style.font
        heading3_font.size = Pt(16)
        heading3_font.bold = True
        heading3_font.color.rgb = RGBColor(
            int(color_scheme['text'][1:3], 16),
            int(color_scheme['text'][3:5], 16),
            int(color_scheme['text'][5:7], 16)
        )
        heading3_style.paragraph_format.space_after = Pt(12)
        heading3_style.paragraph_format.space_before = Pt(15)
        heading3_style.paragraph_format.keep_with_next = True
        
        # Create body text style with enhanced styling
        body_style = doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.size = Pt(12)
        body_font.color.rgb = RGBColor(
            int(color_scheme['text'][1:3], 16),
            int(color_scheme['text'][3:5], 16),
            int(color_scheme['text'][5:7], 16)
        )
        body_style.paragraph_format.space_after = Pt(12)
        body_style.paragraph_format.line_spacing = 1.15
        body_style.paragraph_format.first_line_indent = Pt(0)  # No indent for first line
        
        # Create bullet style with enhanced styling
        bullet_style = doc.styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_font = bullet_style.font
        bullet_font.size = Pt(12)
        bullet_font.color.rgb = RGBColor(
            int(color_scheme['text'][1:3], 16),
            int(color_scheme['text'][3:5], 16),
            int(color_scheme['text'][5:7], 16)
        )
        bullet_style.paragraph_format.space_after = Pt(8)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.first_line_indent = Inches(-0.25)  # Hanging indent
        
        # Create quote style with enhanced styling
        quote_style = doc.styles.add_style('CustomQuote', WD_STYLE_TYPE.PARAGRAPH)
        quote_font = quote_style.font
        quote_font.size = Pt(11)
        quote_font.color.rgb = RGBColor(
            int(color_scheme['quote'][1:3], 16),
            int(color_scheme['quote'][3:5], 16),
            int(color_scheme['quote'][5:7], 16)
        )
        quote_font.italic = True
        quote_style.paragraph_format.space_after = Pt(12)
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.5)
        quote_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Create code style with enhanced styling
        code_style = doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.size = Pt(10)
        code_font.name = 'Courier New'
        code_font.color.rgb = RGBColor(
            int(color_scheme['code'][1:3], 16),
            int(color_scheme['code'][3:5], 16),
            int(color_scheme['code'][5:7], 16)
        )
        code_style.paragraph_format.space_after = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.5)
        # Note: Background shading is complex in python-docx, so we'll just use font styling
        
        # Create insight style for educational content with enhanced styling
        insight_style = doc.styles.add_style('CustomInsight', WD_STYLE_TYPE.PARAGRAPH)
        insight_font = insight_style.font
        insight_font.size = Pt(12)
        insight_font.color.rgb = RGBColor(
            int(color_scheme['insight'][1:3], 16),
            int(color_scheme['insight'][3:5], 16),
            int(color_scheme['insight'][5:7], 16)
        )
        insight_font.bold = True
        insight_style.paragraph_format.space_after = Pt(12)
        insight_style.paragraph_format.line_spacing = 1.15
        insight_style.paragraph_format.first_line_indent = Pt(0)
        
        # Create education style for learning takeaways with enhanced styling
        education_style = doc.styles.add_style('CustomEducation', WD_STYLE_TYPE.PARAGRAPH)
        education_font = education_style.font
        education_font.size = Pt(12)
        education_font.color.rgb = RGBColor(
            int(color_scheme['education'][1:3], 16),
            int(color_scheme['education'][3:5], 16),
            int(color_scheme['education'][5:7], 16)
        )
        education_font.italic = True
        education_style.paragraph_format.space_after = Pt(12)
        education_style.paragraph_format.line_spacing = 1.15
        education_style.paragraph_format.first_line_indent = Pt(0)
    
    def _add_cover_page(self, doc, main_title, topic, color_scheme):
        """Add a professional cover page to the document with enhanced styling and emoji support"""
        # Add title using custom title style
        title_para = doc.add_paragraph(style='CustomTitle')
        self._add_text_with_emoji_support(title_para, main_title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Add document info
        info_para = doc.add_paragraph()
        info_run = info_para.add_run(f"Prepared on: {self._get_current_date()}")
        info_run.font.size = Pt(12)
        info_run.font.color.rgb = RGBColor(
            int(color_scheme['light_text'][1:3], 16),
            int(color_scheme['light_text'][3:5], 16),
            int(color_scheme['light_text'][5:7], 16)
        )
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.paragraph_format.space_after = Pt(20)
        
        # Add page break
        doc.add_page_break()

    def _add_table_of_contents(self, doc, color_scheme):
        """Add a table of contents to the document with enhanced styling and emoji support"""
        # Add TOC heading
        toc_para = doc.add_paragraph(style='CustomHeading1')
        self._add_text_with_emoji_support(toc_para, "📋 Mundarija")
        toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add note about TOC
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("Mundarija hujjatni ochganingizdan so'ng avtomatik yangilanadi")
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        note_run.italic = True
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_para.paragraph_format.space_after = Pt(20)
        
        # Add placeholder for TOC (in a real implementation, this would be more complex)
        placeholder_para = doc.add_paragraph()
        placeholder_run = placeholder_para.add_run("[Mundarija bu yerga avtomatik kiritiladi]")
        placeholder_run.font.size = Pt(11)
        placeholder_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        placeholder_para.paragraph_format.space_after = Pt(30)

    def _add_enhanced_heading(self, doc, text, level=1, colors=None):
        """Add an enhanced heading to the document with emoji support"""
        if level == 1:
            style = 'CustomHeading1'
        elif level == 2:
            style = 'CustomHeading2'
        else:
            style = 'CustomHeading3'
        
        # Add emoji to headings for visual appeal
        emojis = {
            1: "📑 ",
            2: "📋 ",
            3: "📌 "
        }
        enhanced_text = emojis.get(level, "🔹 ") + text
        
        heading_para = doc.add_paragraph(style=style)
        self._add_text_with_emoji_support(heading_para, enhanced_text)
        return heading_para

    def _add_enhanced_paragraph(self, doc, text, colors=None):
        """Add an enhanced paragraph with better formatting"""
        # Check if this is an insight or educational content
        if "💡" in text or "🔑" in text or "📘" in text:
            # Use insight style for key insights
            para = doc.add_paragraph(style='CustomInsight')
            self._add_text_with_emoji_support(para, text)
        elif "🎓" in text or "📚" in text or "📖" in text:
            # Use education style for learning takeaways
            para = doc.add_paragraph(style='CustomEducation')
            self._add_text_with_emoji_support(para, text)
        else:
            # Use normal body style
            para = doc.add_paragraph(style='CustomBody')
            self._add_text_with_emoji_support(para, text)
        return para

    def _add_enhanced_bullet_point(self, doc, text, colors=None):
        """Add an enhanced bullet point to the document with emoji support"""
        para = doc.add_paragraph(style='CustomBullet')
        # Add bullet manually for better control
        run = para.add_run("• ")
        run.bold = True
        self._add_text_with_emoji_support(para, text)
        return para

    def _add_enhanced_numbered_point(self, doc, text, colors=None):
        """Add an enhanced numbered point to the document with emoji support"""
        # For numbered lists, we'll use the built-in numbering
        para = doc.add_paragraph(style='List Number')
        self._add_text_with_emoji_support(para, text)
        return para

    def _add_enhanced_quote(self, doc, text, colors=None):
        """Add an enhanced quote to the document with emoji support"""
        para = doc.add_paragraph(style='CustomQuote')
        self._add_text_with_emoji_support(para, text)
        return para

    def _add_summary_section(self, doc, color_scheme):
        """Add a summary section to the document with enhanced styling and emoji support"""
        # Add page break
        doc.add_page_break()
        
        # Add summary heading
        summary_para = doc.add_paragraph(style='CustomHeading1')
        self._add_text_with_emoji_support(summary_para, "📊 Hujjat Xulosa")
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary points
        summary_points = [
            "✅ Asosiy fikrlar va tushunchalar aniqlandi",
            "✅ Mavzu bo'yicha batafsil tahlil amalga oshirildi",
            "✅ Amaliy tavsiyalar va yechimlar taklif etildi"
        ]
        
        for point in summary_points:
            self._add_enhanced_bullet_point(doc, point)
        
        # Add decorative separator
        doc.add_paragraph()
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run("—" * 50)
        separator_run.font.size = Pt(12)
        separator_run.font.color.rgb = RGBColor(
            int(color_scheme['secondary'][1:3], 16),
            int(color_scheme['secondary'][3:5], 16),
            int(color_scheme['secondary'][5:7], 16)
        )
        separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        separator_para.paragraph_format.space_before = Pt(20)
        separator_para.paragraph_format.space_after = Pt(20)
        
        # Add footer info
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(f"Hujjat tuzildi • {self._get_current_date()}")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(
            int(color_scheme['light_text'][1:3], 16),
            int(color_scheme['light_text'][3:5], 16),
            int(color_scheme['light_text'][5:7], 16)
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.space_before = Pt(10)

    def _add_text_with_emoji_support(self, paragraph, text):
        """Add text to a paragraph with proper emoji and special character support by setting appropriate fonts"""
        # Add the text as a run
        run = paragraph.add_run(text)
        
        # For better emoji and special character support, we need to set appropriate fonts
        # This helps with displaying emojis and special Unicode characters correctly in Word
        try:
            # Set font that supports emojis and special Unicode characters 
            # (try multiple options for cross-platform compatibility)
            unicode_fonts = [
                'Segoe UI Emoji',      # Windows emoji font
                'Segoe UI Symbol',     # Windows symbol font
                'Segoe UI',            # Windows default UI font
                'Apple Color Emoji',   # macOS emoji font
                'Noto Color Emoji',    # Cross-platform emoji font
                'Noto Sans',           # Cross-platform font with good Unicode support
                'Arial Unicode MS',    # Font with extensive Unicode support
                'Tahoma'               # Common Windows font with good Unicode support
            ]
            
            # Try to set the first available font that supports Unicode characters
            for font_name in unicode_fonts:
                try:
                    run.font.name = font_name
                    # For complex scripts and emoji support
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    r.rPr.rFonts.set(qn('w:cs'), font_name)
                    break  # If successful, break the loop
                except Exception:
                    continue  # Try the next font
        except Exception as e:
            # Suppress font warnings as they're not critical and just clutter logs
            # If setting specific font fails, fall back to a more generic approach
            # that should still work for most cases
            # logger.warning(f"Could not set Unicode font: {e}")
            pass
        
        return paragraph
